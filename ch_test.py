# -*- coding: utf-8 -*-
"""
Streamlit 예배 자료 업로드 + Word 저장 + GitHub 임시저장/제출 (+ 성경 JSON 연동)
- '성경 구절' 자료 유형 선택 시: 책/장/절 선택 후 본문 자동 입력
- 성경 JSON은 GitHub 리포의 bsk_json/{book_code}_{chap:03d}.json 에서 로드
"""

# ---------------------------
# 페이지 설정
# ---------------------------
import streamlit as st
st.set_page_config(page_title="설교 자료 업로드", page_icon="🙏", layout="wide")

# ---------------------------
# 표준/서드파티 import
# ---------------------------
import io, os, re, json, uuid, base64, tempfile, requests, hashlib, mimetypes, time
from copy import deepcopy
from typing import List, Dict, Any, Optional
from datetime import date, datetime, timezone
from functools import lru_cache

# python-docx / PIL
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
except Exception:
    st.warning("python-docx가 설치되지 않았습니다. 터미널에서: pip install python-docx")
    Document = None

try:
    from PIL import Image
except Exception:
    st.warning("Pillow가 설치되지 않았습니다. 터미널에서: pip install pillow")
    Image = None

# ---------------------------
# 스타일
# ---------------------------
st.markdown(
    """
    <style>
    .small-note { color:#666; font-size:0.9rem; }
    .section-title { font-weight:700; font-size:1.1rem; margin-top:0.5rem; }
    .landing-card {
        padding: 16px; border: 1px solid #e5e7eb; border-radius: 12px; background: #fff;
        box-shadow: 0 4px 12px rgba(0,0,0,0.04);
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ---------------------------
# 세션 상태 초기화
# ---------------------------
if "materials" not in st.session_state:
    st.session_state.materials: List[Dict[str, Any]] = []
if "preview_idx" not in st.session_state:
    st.session_state.preview_idx = 0
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "role" not in st.session_state:
    st.session_state.role = None
if "user_name" not in st.session_state:
    st.session_state.user_name = ""
if "position" not in st.session_state:
    st.session_state.position = ""
if "can_edit" not in st.session_state:
    st.session_state.can_edit = False
if "worship_date" not in st.session_state:
    st.session_state.worship_date = date.today()
if "submission_id" not in st.session_state:
    st.session_state.submission_id = None

BASE_SERVICES = ["1부", "2부", "3부", "오후예배"]
if "services_options" not in st.session_state:
    st.session_state.services_options = BASE_SERVICES.copy()
if "services_selected" not in st.session_state:
    st.session_state.services_selected: List[str] = []

# ---------------------------
# 성경 JSON 설정 (GitHub 경로/코드/장수)
# ---------------------------
BIBLE_JSON_DIR = st.secrets.get("GITHUB_BIBLE_DIR", "bsk_json")

BOOKS = {
    # OT
    "창세기":"gen","출애굽기":"exo","레위기":"lev","민수기":"num","신명기":"deu",
    "여호수아":"jos","사사기":"jdg","룻기":"rut","사무엘상":"1sa","사무엘하":"2sa",
    "열왕기상":"1ki","열왕기하":"2ki","역대상":"1ch","역대하":"2ch","에스라":"ezr",
    "느헤미야":"neh","에스더":"est","욥기":"job","시편":"psa","잠언":"pro",
    "전도서":"ecc","아가":"sng","이사야":"isa","예레미야":"jer","예레미야애가":"lam",
    "에스겔":"ezk","다니엘":"dan","호세아":"hos","요엘":"jol","아모스":"amo",
    "오바댜":"oba","요나":"jnh","미가":"mic","나훔":"nam","하박국":"hab",
    "스바냐":"zep","학개":"hag","스가랴":"zec","말라기":"mal",
    # NT
    "마태복음":"mat","마가복음":"mrk","누가복음":"luk","요한복음":"jhn","사도행전":"act",
    "로마서":"rom","고린도전서":"1co","고린도후서":"2co","갈라디아서":"gal","에베소서":"eph",
    "빌립보서":"php","골로새서":"col","데살로니가전서":"1th","데살로니가후서":"2th","디모데전서":"1ti",
    "디모데후서":"2ti","디도서":"tit","빌레몬서":"phm","히브리서":"heb","야고보서":"jas",
    "베드로전서":"1pe","베드로후서":"2pe","요한1서":"1jn","요한2서":"2jn","요한3서":"3jn",
    "유다서":"jud","요한계시록":"rev"
}
CHAPTER_COUNT = {
    "창세기":50,"출애굽기":40,"레위기":27,"민수기":36,"신명기":34,"여호수아":24,"사사기":21,"룻기":4,"사무엘상":31,"사무엘하":24,
    "열왕기상":22,"열왕기하":25,"역대상":29,"역대하":36,"에스라":10,"느헤미야":13,"에스더":10,"욥기":42,"시편":150,"잠언":31,
    "전도서":12,"아가":8,"이사야":66,"예레미야":52,"예레미야애가":5,"에스겔":48,"다니엘":12,"호세아":14,"요엘":3,"아모스":9,
    "오바댜":1,"요나":4,"미가":7,"나훔":3,"하박국":3,"스바냐":3,"학개":2,"스가랴":14,"말라기":4,
    "마태복음":28,"마가복음":16,"누가복음":24,"요한복음":21,"사도행전":28,"로마서":16,"고린도전서":16,"고린도후서":13,"갈라디아서":6,"에베소서":6,
    "빌립보서":4,"골로새서":4,"데살로니가전서":5,"데살로니가후서":3,"디모데전서":6,"디모데후서":4,"디도서":3,"빌레몬서":1,"히브리서":13,"야고보서":5,
    "베드로전서":5,"베드로후서":3,"요한1서":5,"요한2서":1,"요한3서":1,"유다서":1,"요한계시록":22
}

# ---------------------------
# 랜딩 (권한/접근)
# ---------------------------
def render_landing():
    st.title("Ch2 설교 자료 업로더")
    st.markdown(
        "<div class='landing-card'>"
        "<b>역할을 선택하고 입장하세요.</b><br>"
        "교역자는 작성/수정이 가능하며, 미디어부는 확인만 가능합니다.<br>"
        "<span class='small-note'>[테스트 안내] 현재는 모든 액세스 코드가 <b>0001</b>이면 입장 가능합니다.</span>"
        "</div>",
        unsafe_allow_html=True
    )
    st.write("")
    with st.form("landing_form"):
        role = st.radio("역할 선택", ["교역자", "미디어부"], horizontal=True)
        c1, c2 = st.columns(2)
        with c1:
            user_name = st.text_input("이름")
        with c2:
            position = st.selectbox(
                "직분 선택",
                ['원로목사', "담임목사", "부목사", '강도사', "전도사", "미디어부"],
                index=2
            )
        access_code = st.text_input("개인 액세스 코드", type="password", placeholder="예) 0001")
        submitted = st.form_submit_button("입장")
    if submitted:
        if access_code == "0001":
            st.session_state.authenticated = True
            st.session_state.role = role
            st.session_state.user_name = user_name.strip()
            st.session_state.position = position
            st.session_state.can_edit = (role == "교역자")
            st.success("입장되었습니다.")
            st.rerun()
        else:
            st.error("액세스 코드가 올바르지 않습니다. (테스트: 0001)")

if not st.session_state.authenticated:
    render_landing()
    st.stop()

# ---------------------------
# 상단 사용자/권한 표시
# ---------------------------
can_edit = st.session_state.get("can_edit", False)
role_badge = "🟢 편집 가능" if can_edit else "🔒 읽기 전용(확인만)"
st.markdown(
    f"**접속자:** {st.session_state.user_name or '이름 미입력'} "
    f"({st.session_state.position or '직분 미선택'}) · "
    f"{st.session_state.role} · {role_badge}"
)

# ---------------------------
# GitHub 유틸
# ---------------------------
def _gh_headers():
    return {
        "Authorization": f"token {st.secrets['GITHUB_TOKEN']}",
        "Accept": "application/vnd.github+json",
    }

def _gh_api_base():
    owner = st.secrets["GITHUB_OWNER"]
    repo = st.secrets["GITHUB_REPO"]
    return f"https://api.github.com/repos/{owner}/{repo}"

def gh_put_bytes(path: str, content_bytes: bytes, message: str):
    api = _gh_api_base()
    url = f"{api}/contents/{path}"
    get = requests.get(url, headers=_gh_headers())
    sha = get.json().get("sha") if get.status_code == 200 else None
    b64 = base64.b64encode(content_bytes).decode("utf-8")
    payload = {
        "message": message,
        "content": b64,
        "branch": st.secrets.get("GITHUB_BRANCH", "main"),
    }
    if sha:
        payload["sha"] = sha
    r = requests.put(url, headers=_gh_headers(), json=payload)
    if r.status_code not in (200, 201):
        raise RuntimeError(f"GitHub 업로드 실패: {r.status_code} {r.text}")
    return r.json()

def gh_get_bytes(path: str) -> bytes:
    api = _gh_api_base()
    url = f"{api}/contents/{path}"
    r = requests.get(url, headers=_gh_headers())
    if r.status_code != 200:
        raise FileNotFoundError(f"GitHub 파일 없음: {path}")
    content = r.json()["content"]
    return base64.b64decode(content)

def gh_list_dir(path: str):
    api = _gh_api_base()
    url = f"{api}/contents/{path}"
    r = requests.get(url, headers=_gh_headers())
    if r.status_code != 200:
        return []
    return r.json()

# ---------------------------
# 자료 유틸
# ---------------------------
def add_material():
    st.session_state.materials.append({
        "id": str(uuid.uuid4()),
        "kind": "성경 구절",
        "files": [],
        "file": None,
        "verse_text": "",
        "description": "",
        "full_text": ""
    })

def remove_material(mid: str):
    st.session_state.materials = [m for m in st.session_state.materials if m["id"] != mid]

def move_material(mid: str, direction: str):
    mats = st.session_state.materials
    idx = next((i for i, m in enumerate(mats) if m["id"] == mid), None)
    if idx is None:
        return
    if direction == "up" and idx > 0:
        mats[idx-1], mats[idx] = mats[idx], mats[idx-1]
        st.session_state.materials = mats
        st.rerun()
    elif direction == "down" and idx < len(mats)-1:
        mats[idx+1], mats[idx] = mats[idx], mats[idx+1]
        st.session_state.materials = mats
        st.rerun()

def add_rich_text(paragraph, text: str):
    if not text:
        return
    pattern = r'(\*\*.*?\*\*|==.*?==)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("==") and part.endswith("=="):
            run = paragraph.add_run(part[2:-2])
            try:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            except Exception:
                pass
        else:
            paragraph.add_run(part)

# ---------------------------
# 성경 JSON 로더 + 피커 위젯
# ---------------------------
@st.cache_data(show_spinner=False, ttl=60*30)
def load_chapter_json_from_github(book_code: str, chap: int) -> Dict[str, Any]:
    json_path = f"{BIBLE_JSON_DIR}/{book_code}_{chap:03d}.json"
    return json.loads(gh_get_bytes(json_path).decode("utf-8"))

def get_book_code(book_name: str) -> str:
    code = BOOKS.get(book_name)
    if not code:
        raise ValueError(f"알 수 없는 책 이름: {book_name}")
    return code

def render_bible_picker(item: Dict[str, Any], disabled: bool):
    st.markdown("**📖 성경 선택**")
    c1, c2, c3 = st.columns([1.4, 0.8, 1.2])
    with c1:
        book_name = st.selectbox(
            "책",
            options=list(BOOKS.keys()),
            key=f"bible_book_{item['id']}",
            disabled=disabled,
        )
    book_code = get_book_code(book_name)
    max_chap = CHAPTER_COUNT[book_name]

    with c2:
        chap = st.number_input("장", min_value=1, max_value=max_chap, step=1,
                               key=f"bible_chap_{item['id']}", disabled=disabled)

    # 절 범위 계산을 위해 해당 장 로드
    data = None
    try:
        data = load_chapter_json_from_github(book_code, int(chap))
        max_verse = len(data.get("verses", [])) or 1
    except Exception as e:
        st.error(f"성경 본문 로드 실패: {e}")
        max_verse = 1

    with c3:
        vcols = st.columns(2)
        with vcols[0]:
            v_from = st.number_input("절(시작)", min_value=1, max_value=max_verse, value=1,
                                     key=f"bible_v_from_{item['id']}", disabled=disabled)
        with vcols[1]:
            v_to = st.number_input("절(끝)", min_value=v_from, max_value=max_verse, value=v_from,
                                   key=f"bible_v_to_{item['id']}", disabled=disabled)

    preview = ""
    if data:
        verses = data.get("verses", [])
        lines = []
        for v in verses:
            vn = v.get("verse")
            if vn is not None and v_from <= vn <= v_to:
                lines.append(f"{book_name} {int(chap)}:{vn} {v.get('text','').strip()}")
        preview = "\n".join(lines)

    st.text_area("미리보기", value=preview, height=140, disabled=True)

    # 🔁 버튼 하나만: 말씀 추가 → 본문 내용(verse_text)에 이어붙이기
    if st.button("📥 말씀 추가", key=f"bible_insert_{item['id']}", disabled=disabled):
        prev = item.get("verse_text", "") or ""
        new_block = preview.strip()
        if new_block:
            item["verse_text"] = (prev + ("\n" if prev else "") + new_block).strip()
            st.success("말씀을 본문 내용에 추가했습니다.")
        else:
            st.warning("추가할 본문이 없습니다. 책/장/절을 확인해 주세요.")

# ---------------------------
# 파일 업로드 보조(메타데이터화)
# ---------------------------
def sanitize_filename(name: str) -> str:
    name = os.path.basename(name or "upload.bin")
    return name.replace("/", "_").replace("\\", "_").strip()

def _sha1(b: bytes) -> str:
    return hashlib.sha1(b).hexdigest()

def upload_streamlit_file_to_github(uploaded_file, dest_dir: str, msg_prefix: str = "[file]") -> dict:
    if uploaded_file is None:
        return {}
    data = uploaded_file.getvalue()
    sha1 = _sha1(data)
    orig_name = getattr(uploaded_file, "name", "upload.bin")
    safe_name = sanitize_filename(orig_name)
    dest_path = f"{dest_dir}/{sha1[:10]}_{safe_name}"
    gh_put_bytes(dest_path, data, message=f"{msg_prefix} upload {safe_name}")
    return {
        "name": orig_name,
        "path": dest_path,
        "size": len(data),
        "content_type": getattr(uploaded_file, "type", mimetypes.guess_type(orig_name)[0]),
        "sha1": sha1,
    }

def materials_upload_and_detach_files(materials: List[Dict[str, Any]], files_dir: str, msg_prefix: str) -> List[Dict[str, Any]]:
    out = []
    for m in materials:
        m2 = deepcopy(m)
        kind = m2.get("kind", "")

        if kind == "이미지":
            metas = []
            files = m2.get("files") or []
            for f in files:
                if hasattr(f, "getvalue"):  # UploadedFile
                    metas.append(upload_streamlit_file_to_github(f, files_dir, msg_prefix))
                elif isinstance(f, dict) and "path" in f:
                    metas.append(f)
            m2["files"] = metas
            m2["file"] = None

        elif kind == "기타 파일":
            f = m2.get("file")
            if hasattr(f, "getvalue"):
                m2["file"] = upload_streamlit_file_to_github(f, files_dir, msg_prefix)
            elif isinstance(f, dict) and "path" in f:
                pass
            else:
                m2["file"] = None

        else:
            if "files" in m2 and not isinstance(m2["files"], list):
                m2["files"] = []
            if "file" in m2 and not isinstance(m2["file"], (dict, type(None))):
                m2["file"] = None

        out.append(m2)
    return out

# ---------------------------
# build_docx
# ---------------------------
def build_docx(worship_date: date, services: List[str], materials: List[Dict[str, Any]],
               user_name: str, position: str, role: str) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx가 설치되지 않았습니다. 'pip install python-docx' 실행 후 다시 시도해주세요.")
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run("설교 자료")
    run.bold = True
    run.font.size = Pt(20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"날짜: {worship_date.strftime('%Y-%m-%d')}\n").bold = True
    meta.add_run("예배 구분: " + (", ".join(services) if services else "(미선택)") + "\n").bold = True
    if user_name or position or role:
        meta.add_run(f"작성자/권한: {user_name or '(미입력)'} ({position or '직분 미선택'}) - {role or '권한 미지정'}").bold = True

    doc.add_paragraph("")
    doc.add_heading("자료 (스토리보드)", level=1)

    if not materials:
        doc.add_paragraph("(추가된 자료가 없습니다)")
    else:
        for idx, item in enumerate(materials, start=1):
            kind = item.get("kind", "")
            verse_text = item.get("verse_text", "") or ""
            description = item.get("description", "") or ""
            full_text = item.get("full_text", "") or ""
            files = item.get("files", []) or []
            single_file = item.get("file")

            doc.add_heading(f"{idx}. {kind}", level=2)

            if kind == "성경 구절":
                if verse_text.strip():
                    for line in verse_text.splitlines():
                        p = doc.add_paragraph()
                        add_rich_text(p, line)
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph("(성경 구절 미입력)")

            elif kind == "이미지":
                if files:
                    for f in files:
                        try:
                            if isinstance(f, dict) and "path" in f:
                                img_bytes = gh_get_bytes(f["path"])
                                _, ext = os.path.splitext(f.get("name") or f["path"])
                                with tempfile.NamedTemporaryFile(delete=False, suffix=ext or ".img") as tmp:
                                    tmp.write(img_bytes)
                                    tmp.flush()
                                    doc.add_picture(tmp.name, width=Inches(5))
                            elif hasattr(f, "getvalue"):
                                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(getattr(f, "name", ""))[1]) as tmp:
                                    tmp.write(f.getvalue())
                                    tmp.flush()
                                    doc.add_picture(tmp.name, width=Inches(5))
                        except Exception:
                            doc.add_paragraph(
                                f"(이미지 삽입 실패) 파일: "
                                f"{(f.get('name') if isinstance(f, dict) else getattr(f, 'name', 'unknown'))}"
                            )
                else:
                    doc.add_paragraph("(이미지 파일 없음)")

            elif kind == "기타 파일":
                if isinstance(single_file, dict) and "name" in single_file:
                    doc.add_paragraph(f"첨부 파일: {single_file['name']} (문서에 직접 삽입되지 않습니다)")
                elif single_file is not None and hasattr(single_file, "getvalue"):
                    doc.add_paragraph(f"첨부 파일: {getattr(single_file, 'name', '파일')} (문서에 직접 삽입되지 않습니다)")
                else:
                    doc.add_paragraph("(첨부 파일 없음)")

            elif kind == "설교 전문":
                if full_text.strip():
                    for line in full_text.splitlines():
                        p = doc.add_paragraph()
                        add_rich_text(p, line)
                else:
                    doc.add_paragraph("(설교 전문 미입력)")

            p = doc.add_paragraph()
            p.add_run("설명(스토리보드): ")
            if description.strip():
                add_rich_text(p, description)
            else:
                p.add_run("(미입력)")

            doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

# ---------------------------
# 제출 직렬화/역직렬화 + 경로
# ---------------------------
def serialize_submission():
    return {
        "worship_date": str(st.session_state.get("worship_date")),
        "services": st.session_state.get("services_selected", []),
        "materials": st.session_state.get("materials", []),
        "user_name": st.session_state.get("user_name"),
        "position": st.session_state.get("position"),
        "role": st.session_state.get("role"),
        "saved_at": datetime.now(timezone.utc).isoformat()
    }

def load_into_session(payload: dict):
    st.session_state.worship_date = date.fromisoformat(payload.get("worship_date"))
    st.session_state.services_selected = payload.get("services", [])
    st.session_state.materials = payload.get("materials", [])

def gh_paths(user_name: str, worship_date: date, submission_id: str = None):
    base = st.secrets.get("GITHUB_BASE_DIR", "worship_submissions")
    d = worship_date.strftime("%Y-%m-%d")
    safe_user = (user_name or "unknown").strip().replace("/", "_")
    sub_id = "draft" if submission_id is None else submission_id
    folder = f"{base}/{d}/{safe_user}/{sub_id}"
    return {
        "folder": folder,
        "files_dir": f"{folder}/files",
        "json": f"{folder}/submission.json",
        "docx": f"{folder}/submission.docx",
    }

# ---------------------------
# ① 날짜/예배 선택
# ---------------------------
st.markdown("<div class='section-title'>① 날짜/예배 선택</div>", unsafe_allow_html=True)
worship_date = st.date_input("예배 날짜", value=st.session_state.worship_date, format="YYYY-MM-DD", disabled=not can_edit)
st.session_state.worship_date = worship_date

c1, c2 = st.columns([2, 1])
with c1:
    st.session_state.services_selected = st.multiselect(
        "예배 구분 선택",
        options=st.session_state.services_options,
        default=st.session_state.services_selected,
        help="해당 날짜에 해당되는 예배를 모두 선택하세요.",
        disabled=not can_edit
    )
with c2:
    new_service = st.text_input("직접 입력", placeholder="예: 청년예배 / 새벽기도", disabled=not can_edit)
    add_new = st.button("추가", disabled=not can_edit)
    if add_new and new_service.strip():
        if new_service not in st.session_state.services_options:
            st.session_state.services_options.append(new_service.strip())
        if new_service not in st.session_state.services_selected:
            st.session_state.services_selected.append(new_service.strip())
        st.rerun()

services = st.session_state.services_selected
st.divider()

# ---------------------------
# ② 자료 추가 (성경/이미지/기타/설교 전문)
# ---------------------------
st.markdown("<div class='section-title'>② 자료 추가 (성경/이미지/기타/설교 전문)</div>", unsafe_allow_html=True)
st.caption("• 설명(스토리보드)에서 **굵게**, ==형광펜== 으로 강조하면 Word에 그대로 반영됩니다.")

add_btn = st.button("+ 자료 추가", disabled=not can_edit)
if add_btn and can_edit:
    add_material()

to_remove: List[str] = []
for i, item in enumerate(st.session_state.materials):
    with st.container(border=True):
        top_cols = st.columns([1.2, 0.2, 0.2, 0.2])
        with top_cols[0]:
            item["kind"] = st.selectbox(
                "자료 유형",
                ["성경 구절", "이미지", "기타 파일", "설교 전문"],
                index=["성경 구절", "이미지", "기타 파일", "설교 전문"].index(item.get("kind", "성경 구절")),
                key=f"kind_{item['id']}",
                disabled=not can_edit
            )
        with top_cols[1]:
            st.write("")
            st.button("▲", key=f"up_{item['id']}", disabled=(not can_edit or i == 0),
                      on_click=move_material, args=(item["id"], "up"))
        with top_cols[2]:
            st.write("")
            st.button("▼", key=f"down_{item['id']}", disabled=(not can_edit or i == len(st.session_state.materials)-1),
                      on_click=move_material, args=(item["id"], "down"))
        with top_cols[3]:
            st.write("")
            if st.button("삭제", key=f"del_{item['id']}", disabled=not can_edit):
                to_remove.append(item["id"])

        if item["kind"] == "성경 구절":
            # ⬇️ 성경(책/장/절) 선택 위젯
            render_bible_picker(item, disabled=not can_edit)
            # ⬇️ 본문 내용(편집 가능)
            item["verse_text"] = st.text_area(
                "본문 내용",  # ← 라벨 변경
                value=item.get("verse_text", ""),
                key=f"verse_{item['id']}",
                height=160,
                disabled=not can_edit
            )
            item["files"], item["file"] = [], None

        elif item["kind"] == "이미지":
            existing = item.get("files") or []
            if existing:
                with st.expander("📷 기존 이미지 보기", expanded=False):
                    names = []
                    for f in existing:
                        if isinstance(f, dict):
                            names.append(f.get("name") or os.path.basename(f.get("path", "")))
                        elif hasattr(f, "name"):
                            names.append(f.name)
                    st.write(", ".join(names) if names else "(목록 없음)")

            new_uploads = st.file_uploader(
                "이미지 업로드 (PNG/JPG) — 여러 장 선택 가능",
                type=["png", "jpg", "jpeg"],
                key=f"files_{item['id']}",
                accept_multiple_files=True,
                disabled=not can_edit
            )
            if new_uploads and len(new_uploads) > 0:
                item["files"] = existing + new_uploads
            else:
                item["files"] = existing
            item["verse_text"] = ""
            item["file"] = None

        elif item["kind"] == "기타 파일":
            existing = item.get("file")
            if existing:
                if isinstance(existing, dict):
                    st.caption(f"기존 첨부: {existing.get('name','(이름 없음)')}")
                elif hasattr(existing, "name"):
                    st.caption(f"기존 첨부: {existing.name}")

            new_one = st.file_uploader(
                "기타 파일 업로드",
                type=None,
                key=f"file_{item['id']}",
                accept_multiple_files=False,
                disabled=not can_edit
            )
            if new_one is not None:
                item["file"] = new_one
            else:
                item["file"] = existing
            item["verse_text"] = ""
            item["files"] = []

        elif item["kind"] == "설교 전문":
            item["full_text"] = st.text_area(
                "설교 전문 입력 (줄바꿈 유지 / **굵게**, ==형광펜== 지원)",
                value=item.get("full_text", ""),
                key=f"full_{item['id']}",
                height=300,
                disabled=not can_edit
            )
            item["verse_text"], item["files"], item["file"] = "", [], None

        item["description"] = st.text_area(
            "설명(스토리보드)",
            value=item.get("description", ""),
            key=f"desc_{item['id']}",
            height=100,
            placeholder="노출 타이밍, 강조 부분 등. **굵게**, ==형광펜== 으로 강조 가능합니다.",
            disabled=not can_edit
        )

if to_remove and can_edit:
    for rid in to_remove:
        remove_material(rid)

st.divider()

# ---------------------------
# ③ Word 파일 생성(로컬 다운로드)
# ---------------------------
st.markdown("<div class='section-title'>③ Word 저장 (로컬 미리 받기)</div>", unsafe_allow_html=True)
col1, col2 = st.columns([1, 2])
with col1:
    do_save = st.button("📄 업로드 하기 (Word 저장)", type="primary", disabled=not can_edit)

if do_save and can_edit:
    try:
        docx_bytes = build_docx(
            worship_date=worship_date,
            services=services,
            materials=st.session_state.materials,
            user_name=st.session_state.user_name,
            position=st.session_state.position,
            role=st.session_state.role
        )
        filename = f"설교자료_{worship_date.strftime('%Y%m%d')}_{'-'.join(services) if services else '미지정'}.docx"
        st.success("Word 파일이 생성되었습니다. 아래 버튼으로 다운로드하세요.")
        st.download_button(
            "⬇️ Word 파일 다운로드",
            data=docx_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        st.error(f"문서 생성 중 오류가 발생했습니다: {e}")

st.divider()

# ---------------------------
# ④ 저장/불러오기/제출 (GitHub)
# ---------------------------
st.markdown("#### 저장/제출")
b1, b2, b3, _ = st.columns([1,1,1,3])
with b1:
    save_draft = st.button("💾 임시 저장", disabled=not can_edit)
with b2:
    load_draft = st.button("↩️ 불러오기", disabled=not can_edit)  # ← 텍스트 단순화
with b3:
    submit_now = st.button("✅ 제출", disabled=not can_edit)

if save_draft and can_edit:
    try:
        p = gh_paths(st.session_state.user_name, worship_date)  # draft
        materials_detached = materials_upload_and_detach_files(
            st.session_state.materials, p["files_dir"], msg_prefix="[draft-files]"
        )
        data = serialize_submission()
        data["materials"] = materials_detached
        gh_put_bytes(
            p["json"],
            json.dumps(data, ensure_ascii=False).encode("utf-8"),
            message=f"[draft] {st.session_state.user_name} {worship_date} 저장"
        )
        st.success("임시 저장되었습니다. (GitHub)")
    except Exception as e:
        st.error(f"임시 저장 실패: {e}")

if load_draft and can_edit:
    try:
        p = gh_paths(st.session_state.user_name, worship_date)  # draft
        draft_bytes = gh_get_bytes(p["json"])
        payload = json.loads(draft_bytes.decode("utf-8"))
        load_into_session(payload)
        st.success("임시 저장본을 불러왔습니다.")
        st.rerun()
    except Exception as e:
        st.error(f"불러오기 실패 또는 저장본 없음: {e}")

if submit_now and can_edit:
    try:
        sub_id = st.session_state.submission_id or datetime.now().strftime("%H%M%S") + "-" + uuid.uuid4().hex[:6]
        st.session_state.submission_id = sub_id
        p = gh_paths(st.session_state.user_name, worship_date, submission_id=sub_id)

        materials_detached = materials_upload_and_detach_files(
            st.session_state.materials, p["files_dir"], msg_prefix="[submit-files]"
        )

        docx_bytes = build_docx(
            worship_date=worship_date,
            services=st.session_state.services_selected,
            materials=st.session_state.materials,
            user_name=st.session_state.user_name,
            position=st.session_state.position,
            role=st.session_state.role
        )

        data = serialize_submission()
        data["status"] = "submitted"
        data["submission_id"] = sub_id
        data["materials"] = materials_detached

        gh_put_bytes(
            p["json"],
            json.dumps(data, ensure_ascii=False).encode("utf-8"),
            message=f"[submit] {st.session_state.user_name} {worship_date} 제출"
        )
        gh_put_bytes(
            p["docx"],
            docx_bytes,
            message=f"[submit-docx] {st.session_state.user_name} {worship_date} DOCX"
        )
        st.success("제출 완료! 미디어부 화면에서 확인 가능합니다.")
    except Exception as e:
        st.error(f"제출 실패: {e}")

st.divider()

# ---------------------------
# ⑤ 미디어부 제출함 (검토/다운로드)
# ---------------------------
if st.session_state.role == "미디어부":
    st.markdown("### 📬 제출함(미디어부) — 날짜별/제출자별 목록")
    base = st.secrets.get("GITHUB_BASE_DIR", "worship_submissions")
    days = gh_list_dir(base)
    if not days:
        st.info("아직 제출된 자료가 없습니다.")
    else:
        day_names = sorted([d["name"] for d in days if d.get("type") == "dir"], reverse=True)
        sel_day = st.selectbox("날짜 선택", options=day_names)
        if sel_day:
            day_dir = f"{base}/{sel_day}"
            users = gh_list_dir(day_dir) or []
            for u in users:
                if u.get("type") != "dir":
                    continue
                with st.expander(f"👤 {u['name']} — {sel_day} 제출물들"):
                    subs = gh_list_dir(u["path"]) or []
                    for s in subs:
                        if s.get("type") != "dir":
                            continue
                        files = gh_list_dir(s["path"]) or []
                        json_item = next((f for f in files if f.get("name") == "submission.json"), None)
                        docx_item = next((f for f in files if f.get("name") == "submission.docx"), None)

                        c1, c2, c3 = st.columns([2, 1, 2])
                        with c1:
                            st.markdown(f"**제출 ID:** {s['name']}")
                        with c2:
                            if json_item:
                                try:
                                    payload = json.loads(gh_get_bytes(json_item["path"]).decode("utf-8"))
                                    info = (
                                        f"- 예배: {', '.join(payload.get('services', [])) or '(미지정)'}\n"
                                        f"- 자료개수: {len(payload.get('materials', []))}\n"
                                        f"- 제출시각(UTC): {payload.get('saved_at','')}\n"
                                    )
                                    st.caption(info)
                                except Exception:
                                    st.caption("메타 로드 실패")
                            else:
                                st.caption("메타 없음")
                        with c3:
                            if docx_item:
                                try:
                                    docx_bytes = gh_get_bytes(docx_item["path"])
                                    st.download_button(
                                        "⬇️ Word 다운로드",
                                        data=docx_bytes,
                                        file_name=f"설교자료_{sel_day}_{u['name']}_{s['name']}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key=f"dl_{sel_day}_{u['name']}_{s['name']}"
                                    )
                                except Exception as e:
                                    st.error(f"다운로드 오류: {e}")
                            else:
                                if json_item and Document is not None:
                                    if st.button("📄 즉석 Word 생성", key=f"mk_{sel_day}_{u['name']}_{s['name']}"):
                                        try:
                                            payload = json.loads(gh_get_bytes(json_item["path"]).decode("utf-8"))
                                            docx_bytes2 = build_docx(
                                                worship_date=date.fromisoformat(sel_day),
                                                services=payload.get("services", []),
                                                materials=payload.get("materials", []),
                                                user_name=payload.get("user_name"),
                                                position=payload.get("position"),
                                                role=payload.get("role")
                                            )
                                            st.download_button(
                                                "⬇️ Word 다운로드(즉석)",
                                                data=docx_bytes2,
                                                file_name=f"설교자료_{sel_day}_{u['name']}_{s['name']}.docx",
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                key=f"dl2_{sel_day}_{u['name']}_{s['name']}"
                                            )
                                        except Exception as e:
                                            st.error(f"생성 오류: {e}")

# ---------------------------
# 풋터
# ---------------------------
st.markdown(
    """
    <hr/>
    <div class='small-note'>
    ⚙️ 이미지 외의 기타 파일은 Word에 직접 삽입되지 않으며, 파일명과 설명이 기록됩니다.<br>
    ✍️ 강조법: **굵게**, ==형광펜== (Word 변환 시 자동 적용)<br>
    🔗 성경 본문은 GitHub의 JSON(bsk_json)에서 로드됩니다.
    </div>
    """,
    unsafe_allow_html=True
)
